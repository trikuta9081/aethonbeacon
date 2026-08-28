from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
PUBLIC = ROOT / "public"
PUBLIC.mkdir(parents=True, exist_ok=True)

PDF_PATH = PUBLIC / "aethon-beacon-app-vision-guide.pdf"
DOCX_PATH = PUBLIC / "aethon-beacon-app-vision-guide.docx"


CONTENT = {
    "title": "NAYIQ App Vision Guide",
    "subtitle": "What the app does, how to use it, and how the daily flow stays calm, private, and guided.",
    "sections": [
        {
            "heading": "1. What this app is for",
            "paragraphs": [
                "NAYIQ is a guidance and redressal app. It is built to help a user say what is happening, understand the situation, choose a useful next step, and move toward calm, clarity, support, or complaint redress without unnecessary wandering.",
                "The app is designed to feel like one guided conversation, not a maze of unrelated screens."
            ],
        },
        {
            "heading": "2. Who can use it",
            "paragraphs": [
                "The app is for students, teachers, parents, doctors, professionals, officers, homemakers, caregivers, retired users, children with supervision, and anyone who wants help thinking clearly about a difficult situation.",
                "The language and flow adapt to the selected role so the user does not have to translate the app into their own life."
            ],
        },
        {
            "heading": "3. The first 30 seconds",
            "bullets": [
                "One clear intake question appears first: What is happening?",
                "The app uses that first line to infer the likely profile, issue type, and urgency.",
                "The app then opens one next destination only: Path, Calm, Help / Redress, Explore, or Guided Support.",
                "The user does not need to hunt across tabs to figure out where to go."
            ],
        },
        {
            "heading": "4. Main tabs and what each one does",
            "table": [
                ["Home", "The first intake page. The user writes or speaks one problem line, then the app routes the next step automatically."],
                ["Path", "Practical guidance for the issue at hand. This tab is where the app explains what the problem means and what to do next."],
                ["Calm", "A quieter layer for slowing the body, reducing pressure, and helping the user settle before making another choice."],
                ["Practice", "Small repeatable exercises that turn a useful idea into one completed action."],
                ["Pattern", "A review of the user's trend data and check-ins so the app can show what is repeating and what needs adjustment."],
                ["Journal", "A place to write, save, and route. The entry becomes a guided note instead of a dead end."],
                ["Explore", "Search across guides, complaint routes, official help, professional help, and community support in one place."],
                ["Community / Messages", "Verified chat, private rooms, and moderated shared posts for useful support and experience sharing."],
                ["Help / Redress", "Complaint routes, office paths, escalation steps, and official redress guidance for serious issues."],
                ["Settings / Profile", "Profile, language, voice, safety controls, reminders, and admin separation."],
                ["Insights / Reports", "Daily and step reports, trend readouts, share/export actions, and retention-supporting summaries."],
            ],
        },
        {
            "heading": "5. How to use the app well",
            "bullets": [
                "Start on Home and type one honest line about the problem.",
                "Choose the profile that best fits your role.",
                "Let the app move you to the next page instead of opening many tabs by hand.",
                "Use Calm when the mind or body feels overwhelmed.",
                "Use Path when you want a clean practical route.",
                "Use Help / Redress when the issue needs complaint filing or escalation.",
                "Use Journal after a step to keep the story and the next move together.",
                "Use Practice and Pattern to keep the flow from becoming decorative."
            ],
        },
        {
            "heading": "6. Reports and privacy",
            "paragraphs": [
                "Verified users receive private reports that are saved locally on the device in this build.",
                "If the user is not verified, the app still generates the report for that session, but it is not saved until phone or email verification is completed.",
                "This behavior is shown to the user at the time of report generation so the privacy rule is clear.",
                "The current release is local-first. That keeps the experience simple while the larger cloud and backup pieces are finished later."
            ],
        },
        {
            "heading": "7. Daily rhythm that keeps the app useful",
            "bullets": [
                "Check in once.",
                "Read the route.",
                "Take one action.",
                "Save the note or report.",
                "Return later for pattern review, practice, or follow-up reminder.",
            ],
        },
        {
            "heading": "8. Safety and trust",
            "paragraphs": [
                "The app is designed to be calm, private, and practical. Urgent situations should move toward SOS, official help, or a real-world professional when needed.",
                "Community spaces are meant to stay moderated and verified. Adult, abusive, or unsafe content should not be allowed to spread there.",
            ],
        },
        {
            "heading": "9. What the user should expect",
            "bullets": [
                "One question first.",
                "One route next.",
                "One useful action after that.",
                "A quieter experience instead of a crowded one.",
                "A product that behaves more like a guide than a maze."
            ],
        },
        {
            "heading": "10. Quick promise",
            "paragraphs": [
                "NAYIQ should feel like: I said my problem, and the app took me somewhere useful.",
            ],
        },
    ],
}


def build_pdf():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="GuideTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=22,
            leading=26,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#102A2D"),
            spaceAfter=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="GuideSubtitle",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=11,
            leading=15,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#4A5B57"),
            spaceAfter=14,
        )
    )
    styles.add(
        ParagraphStyle(
            name="GuideHeading",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=13,
            leading=17,
            textColor=colors.HexColor("#0E6F69"),
            spaceBefore=8,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="GuideBody",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.2,
            leading=14,
            textColor=colors.HexColor("#102A2D"),
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="GuideSmall",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=9,
            leading=12,
            textColor=colors.HexColor("#6F7785"),
            spaceAfter=6,
        )
    )

    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
        title=CONTENT["title"],
        author="Codex",
    )

    story = [
        Paragraph(CONTENT["title"], styles["GuideTitle"]),
        Paragraph(CONTENT["subtitle"], styles["GuideSubtitle"]),
        Paragraph(
            "Front-page access: open this guide only when the front-page button is tapped. It is meant to be a reference, not clutter.",
            styles["GuideSmall"],
        ),
        Spacer(1, 8),
    ]

    for section in CONTENT["sections"]:
        story.append(Paragraph(section["heading"], styles["GuideHeading"]))
        for paragraph in section.get("paragraphs", []):
            story.append(Paragraph(paragraph, styles["GuideBody"]))
        if "bullets" in section:
            bullets = [ListItem(Paragraph(item, styles["GuideBody"])) for item in section["bullets"]]
            story.append(ListFlowable(bullets, bulletType="bullet", start="circle", leftIndent=14))
            story.append(Spacer(1, 4))
        if "table" in section:
            rows = [[Paragraph(cell, styles["GuideBody"]) for cell in row] for row in section["table"]]
            table = Table(rows, colWidths=[28 * mm, 120 * mm], hAlign="LEFT")
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, -1), colors.whitesmoke),
                        ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#102A2D")),
                        ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                        ("FONTSIZE", (0, 0), (-1, -1), 9.5),
                        ("LEADING", (0, 0), (-1, -1), 12),
                        ("ROWBACKGROUNDS", (0, 0), (-1, -1), [colors.HexColor("#FBFEFD"), colors.HexColor("#F4F8F7")]),
                        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D9E8E2")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                        ("TOPPADDING", (0, 0), (-1, -1), 5),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ]
                )
            )
            story.append(table)
            story.append(Spacer(1, 6))

    story.append(Spacer(1, 10))
    story.append(
        Paragraph(
            "Privacy note: this build keeps verified reports private and saved locally on the device. Unverified reports are generated for the current session only and are not saved.",
            styles["GuideSmall"],
        )
    )

    doc.build(story)


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    title = doc.add_paragraph()
    title.alignment = 1
    run = title.add_run(CONTENT["title"])
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = "Arial"

    subtitle = doc.add_paragraph()
    subtitle.alignment = 1
    run = subtitle.add_run(CONTENT["subtitle"])
    run.italic = True
    run.font.size = Pt(10.5)
    run.font.name = "Arial"

    note = doc.add_paragraph()
    note.alignment = 1
    run = note.add_run("Front-page access: open this guide only when the home-screen button is tapped.")
    run.font.size = Pt(9)
    run.font.name = "Arial"

    for section_data in CONTENT["sections"]:
        h = doc.add_paragraph()
        run = h.add_run(section_data["heading"])
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = "Arial"

        for paragraph in section_data.get("paragraphs", []):
            p = doc.add_paragraph(paragraph)
            p.style = doc.styles["Normal"]
            p_format = p.paragraph_format
            p_format.space_after = Pt(4)
            for r in p.runs:
                r.font.size = Pt(10)
                r.font.name = "Arial"

        if "bullets" in section_data:
            for item in section_data["bullets"]:
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(item)
                run.font.size = Pt(10)
                run.font.name = "Arial"
        if "table" in section_data:
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Tab"
            hdr[1].text = "Purpose"
            for row in section_data["table"]:
                cells = table.add_row().cells
                cells[0].text = row[0]
                cells[1].text = row[1]
            doc.add_paragraph()

    final_note = doc.add_paragraph()
    final_note.alignment = 1
    run = final_note.add_run(
        "Privacy note: verified reports are private and saved locally on the device; unverified reports are generated for the session only and are not saved."
    )
    run.font.size = Pt(9)
    run.font.name = "Arial"

    doc.save(DOCX_PATH)


def main():
    build_pdf()
    build_docx()
    print(f"Wrote {PDF_PATH}")
    print(f"Wrote {DOCX_PATH}")


if __name__ == "__main__":
    main()
